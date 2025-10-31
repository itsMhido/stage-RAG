import PyPDF2
import pdfplumber
from pdf2image import convert_from_path
from docx import Document
import os
import tempfile
from .ocr_processor import OCRProcessor

class DocumentProcessor:
    def __init__(self):
        self.ocr = OCRProcessor()
    
    def extract_text_from_pdf(self, pdf_path):
        """Try text extraction first, then OCR if needed"""
        text = ""
        
        # First attempt: Direct text extraction using pdfplumber
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += f"--- Page {page_num + 1} ---\n{page_text.strip()}\n\n"
        except Exception as e:
            print(f"Direct text extraction failed: {e}")
        
        # Second attempt: Try PyPDF2 if pdfplumber failed
        if len(text.strip()) < 50:
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"--- Page {page_num + 1} ---\n{page_text.strip()}\n\n"
            except Exception as e:
                print(f"PyPDF2 extraction failed: {e}")
        
        # If still no substantial text found, use OCR
        if len(text.strip()) < 50:
            print(f"PDF appears to be image-based, using OCR...")
            text = self.ocr_pdf_pages(pdf_path)
        
        return text.strip()
    
    def ocr_pdf_pages(self, pdf_path):
        """Convert PDF pages to images and OCR them"""
        text = ""
        temp_dir = None
        
        try:
            # Create temporary directory for images
            temp_dir = tempfile.mkdtemp()
            
            # Convert PDF to images with high DPI for better OCR
            print("Converting PDF pages to images...")
            images = convert_from_path(pdf_path, dpi=300, fmt='png', output_folder=temp_dir)
            
            for i, image in enumerate(images):
                print(f"Processing page {i+1}/{len(images)} with OCR...")
                
                # Save temporary image
                temp_path = os.path.join(temp_dir, f"page_{i}.png")
                image.save(temp_path, "PNG")
                
                # Extract text using OCR
                page_text = self.ocr.extract_text_from_image(temp_path)
                if page_text and page_text.strip():
                    text += f"--- Page {i+1} ---\n{page_text.strip()}\n\n"
                
                # Clean up temporary file
                try:
                    os.remove(temp_path)
                except:
                    pass
                
        except Exception as e:
            print(f"Error during PDF OCR: {e}")
        finally:
            # Clean up temporary directory
            if temp_dir and os.path.exists(temp_dir):
                try:
                    import shutil
                    shutil.rmtree(temp_dir)
                except:
                    pass
        
        return text
    
    def extract_text_from_docx(self, docx_path):
        """Extract text from Word documents (.docx and .doc)"""
        # First try with python-docx (works for .docx files)
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n"
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            if text.strip():
                return text.strip()
                
        except Exception as e:
            print(f"Error processing DOCX: \"{e}\"")
        
        # For .doc files or if docx failed, try alternative approaches
        return self._extract_text_from_legacy_doc(docx_path)
    
    def _extract_text_from_legacy_doc(self, doc_path):
        """Handle legacy .doc files using alternative methods"""
        print(f"Attempting to process legacy .doc file...")
        
        # Method 1: Try using antiword if available
        text = self._try_antiword(doc_path)
        if text and len(text.strip()) > 10:
            return text.strip()
        
        # Method 2: Try reading as plain text (sometimes works)
        text = self._try_text_extraction(doc_path)
        if text and len(text.strip()) > 10:
            return text.strip()
        
        # Method 3: Convert to image and OCR (last resort)
        print(f"Converting .doc to image for OCR processing...")
        return self._doc_to_ocr(doc_path)
    
    def _try_antiword(self, doc_path):
        """Try using antiword system command if available"""
        try:
            import subprocess
            result = subprocess.run(['antiword', doc_path], 
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
            print(f"Antiword not available or failed: {e}")
        return ""
    
    def _try_text_extraction(self, doc_path):
        """Try to extract readable text from .doc file"""
        try:
            # Sometimes .doc files contain readable text chunks
            with open(doc_path, 'rb') as file:
                content = file.read()
                
            # Try to decode and extract readable text
            text_parts = []
            
            # Look for text patterns in the binary data
            import re
            
            # Try different encodings for text extraction
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'utf-16']:
                try:
                    decoded = content.decode(encoding, errors='ignore')
                    # Extract sequences of printable characters
                    readable_text = re.findall(r'[a-zA-ZÀ-ÿ0-9\s\.,;:!?\-\(\)]{10,}', decoded)
                    text_parts.extend(readable_text)
                except:
                    continue
            
            if text_parts:
                # Clean and join the extracted text
                text = ' '.join(text_parts)
                # Remove excessive whitespace
                text = re.sub(r'\s+', ' ', text)
                return text.strip()
                
        except Exception as e:
            print(f"Text extraction from .doc failed: {e}")
        
        return ""
    
    def _doc_to_ocr(self, doc_path):
        """Convert .doc to image and use OCR (fallback method)"""
        try:
            # This would require LibreOffice or similar to convert .doc to PDF first
            # Then convert PDF to image and OCR
            # For now, return helpful message
            return f"Legacy .doc file detected. Please convert to .docx or PDF format for better text extraction. File: {os.path.basename(doc_path)}"
        except Exception as e:
            return f"Could not process .doc file: {e}"
    
    def extract_text_from_txt(self, txt_path):
        """Read text files with proper encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1', 'windows-1256']
        
        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as file:
                    content = file.read().strip()
                    if content:  # Only return if content is found
                        return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"Error reading file with {encoding}: {e}")
                continue
        
        print(f"Could not read {txt_path} with any encoding")
        return ""
